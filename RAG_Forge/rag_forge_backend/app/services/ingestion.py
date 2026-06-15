"""
RAG Forge – Document Ingestion Service
Handles: parsing (PDF/DOCX/TXT/HTML) → cleaning → chunking → embedding → ChromaDB storage
"""
from __future__ import annotations

import asyncio
import hashlib
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

import aiofiles
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.models import Chunk, Document
from app.services.vector_store import VectorStoreService

import logging
logger = logging.getLogger(__name__)


# ── Text Extraction ──────────────────────────────────────────────
def extract_text_pdf(path: Path) -> tuple[str, int]:
    """Extract text from PDF using PyMuPDF (fitz). Returns (text, page_count)."""
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text("text")
        pages.append(text)
    doc.close()
    return "\n\n".join(pages), len(pages)


def extract_text_docx(path: Path) -> tuple[str, int]:
    """Extract text from DOCX."""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs), 1


def extract_text_txt(path: Path) -> tuple[str, int]:
    """Extract plain text."""
    text = path.read_text(encoding="utf-8", errors="replace")
    return text, 1


def extract_text_html(path: Path) -> tuple[str, int]:
    """Extract text from HTML using BeautifulSoup."""
    from bs4 import BeautifulSoup
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    # Remove scripts/styles
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return text, 1


EXTRACTORS = {
    "pdf":  extract_text_pdf,
    "docx": extract_text_docx,
    "txt":  extract_text_txt,
    "html": extract_text_html,
}


# ── Text Cleaning ─────────────────────────────────────────────────
def clean_text(text: str) -> str:
    """Remove noise: excessive whitespace, null bytes, weird characters."""
    # Remove null bytes
    text = text.replace("\x00", "")
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove excessive blank lines (max 2 in a row)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove excessive spaces within lines
    text = re.sub(r"[ \t]+", " ", text)
    # Strip each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Remove very short lines that are likely noise (page numbers, headers)
    lines = text.split("\n")
    filtered = []
    for line in lines:
        stripped = line.strip()
        # Keep empty lines (paragraph breaks) and substantive lines
        if not stripped or len(stripped) > 3:
            filtered.append(line)
    return "\n".join(filtered).strip()


# ── Smart Chunking ────────────────────────────────────────────────
def chunk_text(
    text: str,
    chunk_size: int = 512,
    overlap: int = 50,
) -> list[dict]:
    """
    Sentence-aware sliding window chunking.
    Returns list of {text, char_count, chunk_index}.
    """
    # Split into sentences (simple heuristic)
    sentence_endings = re.compile(r"(?<=[.!?])\s+")
    sentences = sentence_endings.split(text)

    chunks: list[dict] = []
    current_chunk: list[str] = []
    current_len = 0
    idx = 0

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        sentence_len = len(sentence)

        # If single sentence > chunk_size, split by words
        if sentence_len > chunk_size:
            words = sentence.split()
            word_chunk: list[str] = []
            wlen = 0
            for word in words:
                if wlen + len(word) + 1 > chunk_size and word_chunk:
                    chunk_text_str = " ".join(word_chunk)
                    chunks.append({
                        "text": chunk_text_str,
                        "char_count": len(chunk_text_str),
                        "chunk_index": idx,
                    })
                    idx += 1
                    # Keep overlap words
                    overlap_words = word_chunk[-max(1, overlap // 6):]
                    word_chunk = overlap_words + [word]
                    wlen = sum(len(w) + 1 for w in word_chunk)
                else:
                    word_chunk.append(word)
                    wlen += len(word) + 1
            if word_chunk:
                chunk_text_str = " ".join(word_chunk)
                chunks.append({"text": chunk_text_str, "char_count": len(chunk_text_str), "chunk_index": idx})
                idx += 1
            current_chunk = []
            current_len = 0
            continue

        if current_len + sentence_len + 1 > chunk_size and current_chunk:
            chunk_text_str = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text_str,
                "char_count": len(chunk_text_str),
                "chunk_index": idx,
            })
            idx += 1
            # Overlap: keep last N characters worth of sentences
            overlap_chunks: list[str] = []
            overlap_len = 0
            for prev in reversed(current_chunk):
                if overlap_len + len(prev) > overlap:
                    break
                overlap_chunks.insert(0, prev)
                overlap_len += len(prev) + 1
            current_chunk = overlap_chunks
            current_len = overlap_len

        current_chunk.append(sentence)
        current_len += sentence_len + 1

    if current_chunk:
        chunk_text_str = " ".join(current_chunk)
        chunks.append({"text": chunk_text_str, "char_count": len(chunk_text_str), "chunk_index": idx})

    # Filter empty chunks
    return [c for c in chunks if len(c["text"].strip()) > 20]


# ── Main Ingestion Service ────────────────────────────────────────
class IngestionService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.vector_store = VectorStoreService()

    async def process_document(self, document_id: str) -> None:
        """Full pipeline: parse → clean → chunk → embed → store."""
        doc = await self.db.get(Document, document_id)
        if not doc:
            logger.error(f"Document {document_id} not found")
            return

        try:
            # ── Step 1: Parse ──────────────────────────────────
            await self._update_status(doc, "parsing")
            extractor = EXTRACTORS.get(doc.file_type)
            if not extractor:
                raise ValueError(f"Unsupported file type: {doc.file_type}")

            path = Path(doc.file_path)
            raw_text, page_count = await asyncio.get_event_loop().run_in_executor(
                None, extractor, path
            )
            doc.page_count = page_count

            # ── Step 2: Clean ──────────────────────────────────
            cleaned = clean_text(raw_text)
            if len(cleaned) < 50:
                raise ValueError("Document appears empty after text extraction")

            # ── Step 3: Chunk ──────────────────────────────────
            await self._update_status(doc, "chunking")
            collection = doc.collection
            chunk_dicts = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: chunk_text(cleaned, collection.chunk_size, collection.chunk_overlap),
            )

            # Save chunks to DB
            db_chunks: list[Chunk] = []
            for cd in chunk_dicts:
                chunk = Chunk(
                    document_id=doc.id,
                    chunk_index=cd["chunk_index"],
                    text=cd["text"],
                    char_count=cd["char_count"],
                )
                self.db.add(chunk)
                db_chunks.append(chunk)

            await self.db.flush()  # get IDs
            doc.chunk_count = len(db_chunks)

            # ── Step 4: Embed & Store in ChromaDB ─────────────
            await self._update_status(doc, "embedding")
            texts = [c.text for c in db_chunks]
            ids = [c.id for c in db_chunks]
            metadatas = [
                {
                    "document_id": doc.id,
                    "document_name": doc.original_name,
                    "collection_id": doc.collection_id,
                    "chunk_index": c.chunk_index,
                    "file_type": doc.file_type,
                    "char_count": c.char_count,
                }
                for c in db_chunks
            ]

            collection_name = f"{settings.chroma_collection_prefix}_{doc.collection_id}"
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.vector_store.add_documents(collection_name, ids, texts, metadatas),
            )

            # Update chunk chroma_ids
            for chunk in db_chunks:
                chunk.chroma_id = chunk.id

            doc.embedding_count = len(db_chunks)
            doc.status = "ready"
            doc.indexed_at = datetime.now(timezone.utc)
            await self.db.commit()
            logger.info(f"Document {doc.original_name} indexed: {len(db_chunks)} chunks")

        except Exception as e:
            logger.exception(f"Error processing document {document_id}: {e}")
            doc.status = "error"
            doc.error_message = str(e)[:500]
            await self.db.commit()

    async def _update_status(self, doc: Document, status: str) -> None:
        doc.status = status
        await self.db.commit()

    async def delete_document_vectors(self, document_id: str, collection_id: str) -> None:
        collection_name = f"{settings.chroma_collection_prefix}_{collection_id}"
        await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: self.vector_store.delete_by_document(collection_name, document_id),
        )